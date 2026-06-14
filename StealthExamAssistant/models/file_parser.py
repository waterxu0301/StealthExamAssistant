import os
import base64
from pathlib import Path

class FileParser:
    """文件解析器：支持多种格式的文件解析"""
    
    SUPPORTED_TEXT_EXTENSIONS = {'.txt', '.csv', '.md'}
    SUPPORTED_PDF_EXTENSIONS = {'.pdf'}
    SUPPORTED_DOC_EXTENSIONS = {'.doc', '.docx'}
    SUPPORTED_EXCEL_EXTENSIONS = {'.xls', '.xlsx'}
    SUPPORTED_IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.bmp', '.gif', '.webp'}
    
    @staticmethod
    def get_file_type(file_path):
        """获取文件类型"""
        ext = Path(file_path).suffix.lower()
        
        if ext in FileParser.SUPPORTED_TEXT_EXTENSIONS:
            return 'text'
        elif ext in FileParser.SUPPORTED_PDF_EXTENSIONS:
            return 'pdf'
        elif ext in FileParser.SUPPORTED_DOC_EXTENSIONS:
            return 'doc'
        elif ext in FileParser.SUPPORTED_EXCEL_EXTENSIONS:
            return 'excel'
        elif ext in FileParser.SUPPORTED_IMAGE_EXTENSIONS:
            return 'image'
        else:
            return 'unknown'
            
    @staticmethod
    def parse_file(file_path, use_ocr_for_image=False):
        """解析文件，返回 (text_content, image_base64, file_type)
        
        Args:
            file_path: 文件路径
            use_ocr_for_image: 是否对图片使用 OCR 预处理（节省 Token）
        """
        file_type = FileParser.get_file_type(file_path)
        
        if file_type == 'text':
            return FileParser._parse_text(file_path), None, 'text'
        elif file_type == 'pdf':
            return FileParser._parse_pdf(file_path), None, 'text'
        elif file_type == 'doc':
            return FileParser._parse_doc(file_path), None, 'text'
        elif file_type == 'excel':
            return FileParser._parse_excel(file_path), None, 'text'
        elif file_type == 'image':
            if use_ocr_for_image:
                # 使用 OCR 预处理（节省 Token）
                ocr_text = FileParser._ocr_image(file_path)
                return ocr_text, None, 'text'
            else:
                # 直接返回图片（使用 VLM）
                return None, FileParser._parse_image(file_path), 'image'
        else:
            raise ValueError(f"不支持的文件格式: {Path(file_path).suffix}")
            
    @staticmethod
    def _parse_text(file_path):
        """解析文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            encodings = ['gbk', 'gb2312', 'latin-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except:
                    continue
            raise ValueError(f"无法解析文本文件: {file_path}")
            
    @staticmethod
    def _parse_pdf(file_path):
        """解析 PDF 文件"""
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                        
            return '\n\n'.join(text_parts)
        except ImportError:
            try:
                from PyPDF2 import PdfReader
                
                text_parts = []
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                        
                return '\n\n'.join(text_parts)
            except ImportError:
                raise ImportError("请安装 PDF 解析库: pip install pdfplumber")
                
    @staticmethod
    def _parse_doc(file_path):
        """解析 Word 文档"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
                    
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
                        
            return '\n\n'.join(text_parts)
        except ImportError:
            raise ImportError("请安装 Word 解析库: pip install python-docx")
            
    @staticmethod
    def _parse_excel(file_path):
        """解析 Excel 表格"""
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(file_path, data_only=True)
            text_parts = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"=== 工作表: {sheet_name} ===")
                
                for row in sheet.iter_rows(values_only=True):
                    # 过滤空行
                    row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                    if row_text.strip(' |'):
                        text_parts.append(row_text)
                        
            return '\n\n'.join(text_parts)
        except ImportError:
            # 降级到 xlrd
            try:
                import xlrd
                
                wb = xlrd.open_workbook(file_path)
                text_parts = []
                
                for sheet in wb.sheets():
                    text_parts.append(f"=== 工作表: {sheet.name} ===")
                    
                    for row_idx in range(sheet.nrows):
                        row = [str(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)]
                        row_text = ' | '.join(row)
                        if row_text.strip(' |'):
                            text_parts.append(row_text)
                            
                return '\n\n'.join(text_parts)
            except ImportError:
                raise ImportError("请安装 Excel 解析库: pip install openpyxl")
                
    @staticmethod
    def _ocr_image(file_path):
        """使用 OCR 提取图片中的文字（节省 Token）"""
        try:
            from rapidocr_onnxruntime import RapidOCR
            
            ocr = RapidOCR()
            result, _ = ocr(file_path)
            
            if result is None or len(result) == 0:
                return ""
                
            # 提取文字
            texts = []
            for line in result:
                if line and len(line) >= 2:
                    text = line[1]
                    confidence = line[2]
                    if confidence > 0.5:
                        texts.append(text)
                        
            return '\n'.join(texts)
        except ImportError:
            raise ImportError("请安装 OCR 库: pip install rapidocr_onnxruntime")
        except Exception as e:
            raise ValueError(f"OCR 识别失败: {str(e)}")
            
    @staticmethod
    def _parse_image(file_path):
        """解析图片文件，返回 base64 编码"""
        with open(file_path, 'rb') as f:
            image_data = f.read()
            
        ext = Path(file_path).suffix.lower()
        mime_types = {
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.bmp': 'image/bmp',
            '.gif': 'image/gif',
            '.webp': 'image/webp'
        }
        mime_type = mime_types.get(ext, 'image/png')
        
        base64_data = base64.b64encode(image_data).decode('utf-8')
        
        return base64_data
        
    @staticmethod
    def get_supported_extensions():
        """获取所有支持的文件扩展名"""
        return (
            FileParser.SUPPORTED_TEXT_EXTENSIONS |
            FileParser.SUPPORTED_PDF_EXTENSIONS |
            FileParser.SUPPORTED_DOC_EXTENSIONS |
            FileParser.SUPPORTED_EXCEL_EXTENSIONS |
            FileParser.SUPPORTED_IMAGE_EXTENSIONS
        )
        
    @staticmethod
    def get_file_filter():
        """获取文件对话框的过滤器"""
        exts = FileParser.get_supported_extensions()
        ext_str = ' '.join(f'*{ext}' for ext in exts)
        return f"支持的文件 ({ext_str})"